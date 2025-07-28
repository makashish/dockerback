from PIL import Image
from pdf2image import convert_from_path
from docx import Document
import pytesseract
import os
import uuid
from doc_creator import set_font

# Set Tesseract path (adjust as needed)
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

def process_file(file_path, lang_code, output_folder, language_name='English'):
    output_text = ""

    # OCR processing
    if file_path.lower().endswith(".pdf"):
        pages = convert_from_path(file_path)
        for i, page in enumerate(pages):
            temp_img = f"temp_page_{i}.png"
            page.save(temp_img, "PNG")
            text = pytesseract.image_to_string(Image.open(temp_img), lang=lang_code)
            output_text += text + "\n\n"
            os.remove(temp_img)
    else:
        output_text = pytesseract.image_to_string(Image.open(file_path), lang=lang_code)

    # Create Word document with font applied
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(output_text)
    set_font(run, language_name)  # From doc_creator.py

    # Save DOCX
    filename = f"{uuid.uuid4()}.docx"
    docx_path = os.path.join(output_folder, filename)
    document.save(docx_path)
    return docx_path